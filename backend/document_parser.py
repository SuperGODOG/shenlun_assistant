# 文档解析模块
import os
import logging
from typing import List, Dict, Optional, Tuple
from pathlib import Path

# 文档解析依赖
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available, .docx files will not be supported")

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2 not available, .pdf files will not be supported")

class DocumentParser:
    """文档解析器，支持多种格式的文档解析"""
    
    SUPPORTED_FORMATS = {
        '.txt': 'parse_txt',
        '.md': 'parse_txt',  # Markdown文件按文本处理
        '.docx': 'parse_docx',
        '.pdf': 'parse_pdf'
    }
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def parse_file(self, file_path: str) -> Optional[Dict[str, str]]:
        """解析文件并返回内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            包含标题和内容的字典，失败返回None
        """
        if not os.path.exists(file_path):
            self.logger.error(f"文件不存在: {file_path}")
            return None
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self.SUPPORTED_FORMATS:
            self.logger.error(f"不支持的文件格式: {file_ext}")
            return None
        
        parser_method = getattr(self, self.SUPPORTED_FORMATS[file_ext])
        
        try:
            return parser_method(file_path)
        except Exception as e:
            self.logger.error(f"解析文件失败 {file_path}: {str(e)}")
            return None
    
    def parse_txt(self, file_path: str) -> Dict[str, str]:
        """解析文本文件"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read().strip()
        
        # 使用文件名作为标题
        title = Path(file_path).stem
        
        return {
            'title': title,
            'content': content,
            'format': 'txt'
        }
    
    def parse_docx(self, file_path: str) -> Optional[Dict[str, str]]:
        """解析Word文档"""
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx未安装，无法解析.docx文件")
            return None
        
        try:
            doc = Document(file_path)
            
            # 提取标题（第一个段落或文件名）
            title = Path(file_path).stem
            if doc.paragraphs and doc.paragraphs[0].text.strip():
                first_para = doc.paragraphs[0].text.strip()
                if len(first_para) < 100:  # 如果第一段较短，可能是标题
                    title = first_para
            
            # 提取所有段落内容
            content_parts = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)
            
            content = '\n'.join(content_parts)
            
            return {
                'title': title,
                'content': content,
                'format': 'docx'
            }
            
        except Exception as e:
            self.logger.error(f"解析Word文档失败: {str(e)}")
            return None
    
    def parse_pdf(self, file_path: str) -> Optional[Dict[str, str]]:
        """解析PDF文档"""
        if not PDF_AVAILABLE:
            self.logger.error("PyPDF2未安装，无法解析.pdf文件")
            return None
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                # 提取所有页面的文本
                content_parts = []
                for page in pdf_reader.pages:
                    text = page.extract_text().strip()
                    if text:
                        content_parts.append(text)
                
                content = '\n'.join(content_parts)
                
                # 使用文件名作为标题
                title = Path(file_path).stem
                
                return {
                    'title': title,
                    'content': content,
                    'format': 'pdf'
                }
                
        except Exception as e:
            self.logger.error(f"解析PDF文档失败: {str(e)}")
            return None
    
    def parse_content_from_text(self, content: str, filename: str = "untitled") -> Dict[str, str]:
        """从文本内容直接解析
        
        Args:
            content: 文本内容
            filename: 文件名（用作标题）
            
        Returns:
            包含标题和内容的字典
        """
        # 尝试从内容中提取标题
        lines = content.strip().split('\n')
        title = filename
        
        if lines:
            first_line = lines[0].strip()
            # 如果第一行较短且不包含太多标点，可能是标题
            if len(first_line) < 100 and first_line.count('.') < 3:
                title = first_line
        
        return {
            'title': title,
            'content': content.strip(),
            'format': 'text'
        }
    
    def batch_parse_directory(self, directory: str, recursive: bool = True) -> List[Dict[str, str]]:
        """批量解析目录中的文档
        
        Args:
            directory: 目录路径
            recursive: 是否递归解析子目录
            
        Returns:
            解析结果列表
        """
        results = []
        
        if not os.path.exists(directory):
            self.logger.error(f"目录不存在: {directory}")
            return results
        
        # 获取所有支持的文件
        files_to_parse = []
        
        if recursive:
            for root, dirs, files in os.walk(directory):
                for file in files:
                    file_path = os.path.join(root, file)
                    if Path(file_path).suffix.lower() in self.SUPPORTED_FORMATS:
                        files_to_parse.append(file_path)
        else:
            for file in os.listdir(directory):
                file_path = os.path.join(directory, file)
                if os.path.isfile(file_path) and Path(file_path).suffix.lower() in self.SUPPORTED_FORMATS:
                    files_to_parse.append(file_path)
        
        # 解析文件
        for file_path in files_to_parse:
            result = self.parse_file(file_path)
            if result:
                result['source_path'] = file_path
                results.append(result)
                self.logger.info(f"成功解析文件: {file_path}")
        
        self.logger.info(f"批量解析完成，共处理 {len(results)} 个文件")
        return results
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式列表"""
        supported = ['.txt', '.md']  # 基础支持
        
        if DOCX_AVAILABLE:
            supported.append('.docx')
        
        if PDF_AVAILABLE:
            supported.append('.pdf')
        
        return supported
    
    def validate_file_size(self, file_path: str, max_size_mb: int = 10) -> bool:
        """验证文件大小
        
        Args:
            file_path: 文件路径
            max_size_mb: 最大文件大小（MB）
            
        Returns:
            文件大小是否符合要求
        """
        try:
            file_size = os.path.getsize(file_path)
            max_size_bytes = max_size_mb * 1024 * 1024
            
            if file_size > max_size_bytes:
                self.logger.warning(f"文件过大: {file_path} ({file_size / 1024 / 1024:.2f}MB > {max_size_mb}MB)")
                return False
            
            return True
        except Exception as e:
            self.logger.error(f"检查文件大小失败: {str(e)}")
            return False

# 全局文档解析器实例
_document_parser = None

def get_document_parser() -> DocumentParser:
    """获取文档解析器实例"""
    global _document_parser
    if _document_parser is None:
        _document_parser = DocumentParser()
    return _document_parser

def parse_document(file_path: str) -> Optional[Dict[str, str]]:
    """便捷函数：解析单个文档"""
    parser = get_document_parser()
    return parser.parse_file(file_path)

def parse_documents_batch(directory: str, recursive: bool = True) -> List[Dict[str, str]]:
    """便捷函数：批量解析文档"""
    parser = get_document_parser()
    return parser.batch_parse_directory(directory, recursive)